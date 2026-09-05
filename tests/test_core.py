from __future__ import annotations

import csv
import sqlite3
import unittest
from io import BytesIO, StringIO
from pathlib import Path
from unittest.mock import Mock, patch
from uuid import uuid4

from docx import Document
from pydantic import ValidationError

import core


def make_bundle() -> core.ApplicationBundle:
    return core.ApplicationBundle(
        analysis={
            "match_score": 78,
            "matched_skills": ["Python", "APIs"],
            "missing_skills": ["Kubernetes"],
            "seniority_fit": "matched",
            "notes": "Strong backend overlap with one infrastructure gap.",
            "requirements": [
                {
                    "requirement": "Build Python APIs",
                    "importance": "must-have",
                    "evidence": "Built FastAPI services for a production project.",
                    "fit": "matched",
                    "confidence": 95,
                    "action": "Lead with the FastAPI project.",
                }
            ],
        },
        cover_letter=(
            "I built Python API services and shipped them with a small engineering team. "
            "That experience maps directly to this backend role.\n\n"
            "My resume also shows practical testing and delivery experience. I would welcome "
            "the opportunity to discuss how that background can support the team."
        ),
        suggested_bullets=[
            {
                "original_evidence": "Built FastAPI services for a production project.",
                "suggested_bullet": "Built and shipped production FastAPI services.",
                "reason": "Uses the role's API terminology without adding facts.",
            }
        ],
        interview_questions=[
            {
                "question": f"Question {number}?",
                "category": "technical" if number % 2 else "behavioural",
                "tip": "Use a specific example from the resume.",
            }
            for number in range(1, 5)
        ],
    )


class CoreTestCase(unittest.TestCase):
    def setUp(self) -> None:
        self.original_database = core.DATABASE_FILE
        self.original_tracker = core.TRACKER_FILE
        self.original_url = core.DATABASE_URL
        # Pin every test to a throwaway SQLite file. Without this, a developer
        # with DATABASE_URL configured would run the suite against the real
        # hosted database, writing test rows into it.
        core.DATABASE_URL = ""
        identifier = uuid4().hex
        test_directory = Path(__file__).parent
        core.DATABASE_FILE = str(test_directory / f".test-{identifier}.db")
        core.TRACKER_FILE = str(test_directory / f".test-{identifier}.csv")

    def tearDown(self) -> None:
        generated = [
            Path(core.DATABASE_FILE),
            Path(f"{core.DATABASE_FILE}-shm"),
            Path(f"{core.DATABASE_FILE}-wal"),
            Path(core.TRACKER_FILE),
        ]
        core.DATABASE_FILE = self.original_database
        core.TRACKER_FILE = self.original_tracker
        core.DATABASE_URL = self.original_url
        for path in generated:
            path.unlink(missing_ok=True)

    def save_sample(self, bundle: core.ApplicationBundle | None = None) -> str:
        return core.save_application(
            company="Acme",
            role="Backend Engineer",
            job_text="Build Python services and operate APIs.",
            bundle=bundle or make_bundle(),
            model="gpt-4o-mini",
            usage={"input_tokens": 100, "output_tokens": 50, "total_tokens": 150},
            resume_name="resume.docx",
            resume_text="Built FastAPI services for a production project.",
            job_url="https://example.com/jobs/1",
            source="Referral",
            location="Remote",
        )

    def test_extract_txt_resume_normalizes_text(self) -> None:
        text = core.extract_resume_text(b"Name\r\n\r\n  Python   Developer  ", "resume.txt")
        self.assertEqual(text, "Name\nPython Developer")

    def test_extract_resume_rejects_unknown_format(self) -> None:
        with self.assertRaisesRegex(core.AssistantError, "Unsupported resume format"):
            core.extract_resume_text(b"hello", "resume.rtf")

    def test_docx_parser_reads_paragraphs_and_tables(self) -> None:
        document = Document()
        document.add_paragraph("Backend engineer")
        table = document.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "Python"
        table.cell(0, 1).text = "FastAPI"
        output = BytesIO()
        document.save(output)

        text = core.extract_resume_text(output.getvalue(), "resume.docx")

        self.assertIn("Backend engineer", text)
        self.assertIn("Python | FastAPI", text)

    def test_schema_rejects_invalid_fit_score(self) -> None:
        payload = make_bundle().model_dump()
        payload["analysis"]["match_score"] = 101
        with self.assertRaises(ValidationError):
            core.ApplicationBundle.model_validate(payload)

    def test_database_lifecycle_and_status_history(self) -> None:
        application_id = self.save_sample()
        saved = core.get_application(application_id)
        self.assertIsNotNone(saved)
        self.assertEqual(saved["company"], "Acme")
        self.assertEqual(saved["status"], "drafted")
        self.assertEqual(saved["analysis"]["match_score"], 78)

        core.update_status(application_id, "applied", "Submitted through careers page")
        core.update_artifacts(application_id, cover_letter="Updated truthful cover letter. " * 5)

        updated = core.get_application(application_id)
        self.assertIsNotNone(updated)
        self.assertEqual(updated["status"], "applied")
        self.assertTrue(updated["cover_letter"].startswith("Updated truthful"))
        history = core.get_status_history(application_id)
        self.assertEqual(history[0]["to_status"], "applied")
        self.assertEqual(len(history), 2)

        core.delete_application(application_id)
        self.assertIsNone(core.get_application(application_id))

    def test_legacy_csv_is_migrated_once(self) -> None:
        csv_path = Path(core.TRACKER_FILE)
        with csv_path.open("w", newline="", encoding="utf-8") as handle:
            writer = csv.DictWriter(
                handle,
                fieldnames=["date", "company", "role", "match_score", "status", "notes"],
            )
            writer.writeheader()
            writer.writerow(
                {
                    "date": "2026-01-10",
                    "company": "Legacy Co",
                    "role": "Engineer",
                    "match_score": "82",
                    "status": "interview",
                    "notes": "Imported row",
                }
            )

        core.init_database()
        core.init_database()
        applications = core.list_applications()

        self.assertEqual(len(applications), 1)
        self.assertEqual(applications[0]["company"], "Legacy Co")
        self.assertEqual(applications[0]["status"], "interviewing")

    def test_csv_export_prevents_formula_injection(self) -> None:
        application_id = self.save_sample()
        application = core.get_application(application_id)
        self.assertIsNotNone(application)
        application["company"] = "=HYPERLINK(\"https://example.com\")"

        exported = core.export_applications_csv([application])

        self.assertIn("'=HYPERLINK", exported)

    def test_docx_export_contains_current_edit(self) -> None:
        text = "Current edited paragraph. " * 8
        data = core.build_cover_letter_docx(text, "Acme", "Backend Engineer")
        document = Document(BytesIO(data))
        extracted = "\n".join(paragraph.text for paragraph in document.paragraphs)
        self.assertIn("Backend Engineer — Acme", extracted)
        self.assertIn("Current edited paragraph", extracted)

    def test_pipeline_uses_structured_output_and_disables_storage(self) -> None:
        bundle = make_bundle()

        class Usage:
            input_tokens = 120
            output_tokens = 80
            total_tokens = 200

        class Response:
            output_parsed = bundle
            usage = Usage()

        class Responses:
            def __init__(self) -> None:
                self.arguments = {}

            def parse(self, **kwargs):
                self.arguments = kwargs
                return Response()

        class Client:
            def __init__(self) -> None:
                self.responses = Responses()

        client = Client()
        result = core.run_pipeline(
            client,
            "Built FastAPI services for a production project.",
            "Build Python services and operate APIs.",
            "Acme",
            "Backend Engineer",
            resume_name="resume.docx",
        )

        self.assertIs(client.responses.arguments["text_format"], core.ApplicationBundle)
        self.assertFalse(client.responses.arguments["store"])
        self.assertEqual(result["usage"]["total_tokens"], 200)
        self.assertIsNotNone(core.get_application(result["application_id"]))

    def test_job_url_validation_rejects_non_http_urls(self) -> None:
        for url in ("javascript:alert(1)", "example.com/job", "file:///tmp/job.txt"):
            with self.subTest(url=url), self.assertRaisesRegex(core.AssistantError, "http"):
                core.validate_job_url(url)

    def test_malformed_http_urls_raise_friendly_errors(self) -> None:
        for url in ("https://[broken", "https://example.com:bad/job", "https:// /job"):
            with self.subTest(url=url), self.assertRaises(core.AssistantError):
                core.validate_job_url(url)

    def test_cover_letter_save_and_export_preserve_paragraphs(self) -> None:
        application_id = self.save_sample()
        text = "Dear team,\r\n\r\nI build Python services.\r\n\r\nRegards,\r\nCandidate"
        core.update_artifacts(application_id, cover_letter=text)
        saved = core.get_application(application_id)["cover_letter"]
        self.assertEqual(saved, text.replace("\r\n", "\n"))
        document = Document(BytesIO(core.build_cover_letter_docx(saved, "Acme", "Engineer")))
        self.assertEqual(
            [paragraph.text for paragraph in document.paragraphs][1:],
            ["Dear team,", "I build Python services.", "Regards,\nCandidate"],
        )

    def test_non_utf8_job_file_raises_friendly_error(self) -> None:
        path = Path(core.TRACKER_FILE).with_suffix(".txt")
        try:
            path.write_bytes(b"Developer \xff")
            with self.assertRaisesRegex(core.AssistantError, "UTF-8"):
                core.load_job_posting(str(path))
        finally:
            path.unlink(missing_ok=True)

    def test_database_read_and_delete_errors_are_friendly(self) -> None:
        core.init_database()
        operations = [
            lambda: core.get_application("missing"),
            lambda: core.delete_application("missing"),
            lambda: core.get_status_history("missing"),
            core.dashboard_stats,
        ]
        for operation in operations:
            with self.subTest(operation=operation), patch.object(
                core, "_open_sqlite", side_effect=sqlite3.OperationalError("unavailable")
            ), self.assertRaises(core.AssistantError):
                operation()

    def test_pipeline_checks_storage_before_spending_api_tokens(self) -> None:
        client = Mock()
        with patch.object(
            core, "init_database", side_effect=core.AssistantError("offline")
        ), self.assertRaisesRegex(core.AssistantError, "offline"):
            core.run_pipeline(client, "Resume text", "Job text", "Acme", "Engineer")
        client.responses.parse.assert_not_called()

    def test_inline_status_edits_resolve_to_the_edited_application(self) -> None:
        rows = [
            {"id": "first", "status": "drafted"},
            {"id": "second", "status": "applied"},
        ]
        # The editable table reports edits keyed by row position.
        self.assertEqual(
            core.resolve_status_edits({0: {"status": "applied"}}, rows),
            [("first", "applied")],
        )
        # Re-selecting the value a row already has is not a change.
        self.assertEqual(core.resolve_status_edits({1: {"status": "applied"}}, rows), [])
        # Edits to other columns are ignored.
        self.assertEqual(core.resolve_status_edits({0: {"notes": "x"}}, rows), [])

    def test_inline_status_edits_skip_positions_that_no_longer_resolve(self) -> None:
        """A stale row position must never write a status onto the wrong application."""
        rows = [{"id": "only", "status": "drafted"}]
        for edits in ({5: {"status": "offer"}}, {"bad": {"status": "offer"}}, None, "x"):
            with self.subTest(edits=edits):
                self.assertEqual(core.resolve_status_edits(edits, rows), [])

    def test_invalid_legacy_csv_encoding_is_friendly_and_retryable(self) -> None:
        path = Path(core.TRACKER_FILE)
        path.write_bytes(b"company,role\nAcme,\xff\n")
        with self.assertRaisesRegex(core.AssistantError, "legacy tracker"):
            core.init_database()
        path.write_text("company,role\nAcme,Engineer\n", encoding="utf-8")
        core.init_database()
        self.assertEqual(len(core.list_applications()), 1)

    def test_export_includes_all_rows_beyond_display_limit_and_respects_filters(self) -> None:
        with Path(core.TRACKER_FILE).open("w", newline="", encoding="utf-8") as handle:
            writer = csv.writer(handle)
            writer.writerow(["company", "role", "status"])
            writer.writerows(["Acme", "Engineer", "drafted"] for _ in range(2001))
            writer.writerow(["Other", "Designer", "applied"])
        self.assertEqual(len(core.list_applications()), 500)
        self.assertEqual(len(list(csv.DictReader(StringIO(core.export_applications_csv())))), 2002)
        filtered = list(csv.DictReader(StringIO(
            core.export_applications_csv(search="Acme", status="drafted")
        )))
        self.assertEqual(len(filtered), 2001)
        self.assertTrue(all(row["company"] == "Acme" for row in filtered))

    def test_csv_export_neutralizes_formulas_after_whitespace(self) -> None:
        for value in ("\n=1+1", "  =1+1", "\t@SUM(1)"):
            with self.subTest(value=value):
                exported = core.export_applications_csv([{"company": value}])
                row = next(csv.DictReader(StringIO(exported)))
                self.assertEqual(row["company"], "'" + value)

    def test_failed_save_rolls_back_application_and_event(self) -> None:
        core.init_database()
        original_execute = core._Connection.execute

        def fail_event(connection, sql, parameters=()):
            if "INSERT INTO application_events" in sql:
                raise sqlite3.OperationalError("simulated event write failure")
            return original_execute(connection, sql, parameters)

        with patch.object(core._Connection, "execute", fail_event), self.assertRaises(core.AssistantError):
            self.save_sample()
        self.assertEqual(core.list_applications(), [])


class PostgresDialectTestCase(unittest.TestCase):
    """The Postgres backend reuses SQLite-flavoured SQL, so translation matters."""

    def test_placeholders_become_pyformat(self) -> None:
        translated = core._to_postgres(
            "INSERT INTO metadata(key, value) VALUES('legacy_csv_migrated', ?)"
        )
        self.assertNotIn("?", translated)
        self.assertIn("VALUES('legacy_csv_migrated', %s)", translated)

    def test_like_becomes_ilike_to_stay_case_insensitive(self) -> None:
        translated = core._to_postgres(
            "SELECT * FROM applications WHERE (company LIKE ? OR role LIKE ?) LIMIT ?"
        )
        self.assertEqual(
            translated,
            "SELECT * FROM applications WHERE (company ILIKE %s OR role ILIKE %s) LIMIT %s",
        )

    def test_like_substrings_are_not_rewritten(self) -> None:
        self.assertEqual(core._to_postgres("SELECT unliked FROM t"), "SELECT unliked FROM t")

    def test_no_literal_percent_in_translated_statements(self) -> None:
        """A stray % would be read as a psycopg placeholder and break the query."""
        for schema in (core.SCHEMA_SQLITE, core.SCHEMA_POSTGRES):
            self.assertNotIn("%", schema)

    def test_postgres_schema_replaces_sqlite_autoincrement(self) -> None:
        self.assertIn("AUTOINCREMENT", core.SCHEMA_SQLITE)
        self.assertNotIn("AUTOINCREMENT", core.SCHEMA_POSTGRES)
        self.assertIn("GENERATED BY DEFAULT AS IDENTITY", core.SCHEMA_POSTGRES)

    def test_engine_selection_follows_database_url(self) -> None:
        original = core.DATABASE_URL
        try:
            core.DATABASE_URL = ""
            self.assertFalse(core.use_postgres())
            core.DATABASE_URL = "postgresql://user:pw@host/db"
            self.assertTrue(core.use_postgres())
        finally:
            core.DATABASE_URL = original

    @unittest.skipIf(core.psycopg is None, "Postgres driver is not installed")
    def test_postgres_connection_errors_are_friendly(self) -> None:
        pool = Mock()
        pool.connection.side_effect = core.psycopg.OperationalError("connection lost")
        with patch.object(core, "DATABASE_URL", "postgresql://test.invalid/test"), patch.object(
            core, "_get_pool", return_value=pool
        ), self.assertRaises(core.AssistantError), core._database():
            self.fail("An unavailable database must not yield a connection")


if __name__ == "__main__":
    unittest.main()
